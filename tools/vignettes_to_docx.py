# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "anthropic>=0.40",
#     "python-docx>=1.1",
#     "pillow>=10.4",
# ]
# ///
"""Dev utility: transcribe image vignettes with the Claude vision API into a .docx.

Standalone (PEP 723 inline deps) — not part of the MediSnap app. Handy for
turning a folder of photographed notes into text for testing.

Usage:
    export ANTHROPIC_API_KEY=sk-ant-...
    uv run tools/vignettes_to_docx.py [SRC_DIR] [OUT.docx]

Defaults:
    SRC_DIR  = ./vignettes
    OUT.docx = <SRC_DIR>/vignettes.docx
"""

import base64
import io
import sys
from pathlib import Path

import anthropic
from docx import Document
from PIL import Image

MODEL = "claude-opus-4-8"
MAX_EDGE = 1568  # Claude downsizes above this anyway; keeps us under the 5MB/image limit
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
PROMPT = (
    "This is a photo of a written/printed vignette. Transcribe ALL of its text "
    "exactly as it appears, preserving paragraph breaks and reading order. "
    "Output only the transcription — no commentary, headings, or markdown fences."
)


def encode_resized(path: Path) -> str:
    """Downscale the long edge to MAX_EDGE and return base64 JPEG."""
    img = Image.open(path).convert("RGB")
    if max(img.size) > MAX_EDGE:
        scale = MAX_EDGE / max(img.size)
        img = img.resize((round(img.width * scale), round(img.height * scale)))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


def transcribe(client: anthropic.Anthropic, path: Path) -> str:
    data = encode_resized(path)
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": data,
                        },
                    },
                    {"type": "text", "text": PROMPT},
                ],
            }
        ],
    )
    return "".join(b.text for b in response.content if b.type == "text").strip()


def main() -> int:
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("vignettes")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src / "vignettes.docx"

    if not src.is_dir():
        print(f"Source directory not found: {src}", file=sys.stderr)
        return 1

    images = sorted(p for p in src.iterdir() if p.suffix.lower() in IMAGE_EXTS)
    if not images:
        print(f"No images found in {src}", file=sys.stderr)
        return 1

    client = anthropic.Anthropic()  # ANTHROPIC_API_KEY from env
    doc = Document()
    doc.add_heading("Vignettes", level=0)

    for i, path in enumerate(images, 1):
        print(f"[{i}/{len(images)}] {path.name} …", flush=True)
        try:
            text = transcribe(client, path)
        except Exception as e:  # noqa: BLE001 - record the failure, keep going
            text = f"[ERROR transcribing {path.name}: {e}]"
            print(f"    failed: {e}", file=sys.stderr)

        doc.add_heading(f"Vignette {i} — {path.name}", level=1)
        for para in text.split("\n"):
            doc.add_paragraph(para)
        if i < len(images):
            doc.add_page_break()

    doc.save(out)
    print(f"\nSaved {len(images)} vignettes to {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
